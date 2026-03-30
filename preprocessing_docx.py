from pathlib import Path
from collections import Counter
import json
import re
import os
import win32com.client

from docx import Document


def normalize_text(text: str) -> str:
    if not text:
        return ""

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def remove_consecutive_duplicates(lines: list[str]) -> list[str]:
    cleaned = []
    prev = None

    for line in lines:
        line = normalize_text(line)
        if not line:
            continue

        if line != prev:
            cleaned.append(line)
        prev = line

    return cleaned


def remove_repeated_lines(lines: list[str], min_repeat: int = 3, max_length: int = 80) -> list[str]:
    normalized_lines = [normalize_text(line) for line in lines if normalize_text(line)]
    counter = Counter(normalized_lines)

    cleaned = []
    for line in normalized_lines:
        if len(line) <= max_length and counter[line] >= min_repeat:
            continue
        cleaned.append(line)

    return cleaned


def table_to_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [normalize_text(cell.text) for cell in row.cells]
        cells = [c for c in cells if c]
        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows).strip()

def convert_doc_to_docx(doc_path: str) -> str:
    doc_path = Path(doc_path).resolve()
    
    if not doc_path.exists():
        raise FileNotFoundError(f"File not found: {doc_path}")
    
    new_path = doc_path.with_suffix(".docx")

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        doc = word.Documents.Open(str(doc_path))
        doc.SaveAs(str(new_path), FileFormat=16)
        doc.Close()
    finally:
        word.Quit()

    return str(new_path)

def ensure_docx_format(file_path: str) -> str:
    path = Path(file_path).resolve()

    if path.suffix.lower() == ".docx":
        return str(path)

    elif path.suffix.lower() == ".doc":
        print(f"[INFO] Convert .doc → .docx: {path.name}")
        new_path = convert_doc_to_docx(str(path))
        return new_path

    else:
        raise ValueError(f"지원하지 않는 파일 형식: {path.suffix}")

def extract_docx_blocks(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    blocks = []

    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"

        block_type = "heading" if "Heading" in style_name else "paragraph"

        blocks.append({
            "type": block_type,
            "style": style_name,
            "text": text
        })

    for table in doc.tables:
        table_text = table_to_text(table)
        if table_text:
            blocks.append({
                "type": "table",
                "style": "Table",
                "text": table_text
            })

    return blocks


def preprocess_docx_blocks(blocks: list[dict]) -> list[dict]:
    raw_lines = [block["text"] for block in blocks]
    lines_no_dup = remove_consecutive_duplicates(raw_lines)

    lines_no_repeat = remove_repeated_lines(lines_no_dup)

    cleaned_blocks = []
    used = set()

    for block in blocks:
        text = normalize_text(block["text"])
        if text in lines_no_repeat and text not in used:
            cleaned_blocks.append({
                "type": block["type"],
                "style": block["style"],
                "text": text
            })
            used.add(text)

    return cleaned_blocks


def save_preprocessed_docx(docx_path: str, output_path: str):
    docx_path = ensure_docx_format(docx_path)
    blocks = extract_docx_blocks(docx_path)

    cleaned_blocks = preprocess_docx_blocks(blocks)

    result = {
        "source_file": str(Path(docx_path).name),
        "num_blocks": len(cleaned_blocks),
        "blocks": cleaned_blocks
    }

    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    return result


if __name__ == "__main__":
    docx_path = r"input\docx\test_docx1.doc"
    output_path = r"output\json\preprocessed_docx1.json"

    result = save_preprocessed_docx(docx_path, output_path)
    print(f"Saved successfully: {output_path}")
    print(f"Preprocessed blocks: {result['num_blocks']}")