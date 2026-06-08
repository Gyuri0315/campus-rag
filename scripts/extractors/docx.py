from __future__ import annotations

from pathlib import Path

from scripts.extractors.common import normalize_text


def extract_docx_like_blocks(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    work_path = path
    if suffix == ".doc":
        try:
            import win32com.client  # type: ignore
        except Exception as exc:
            raise RuntimeError(f".doc conversion requires pywin32: {exc}") from exc

        converted = path.with_suffix(".docx")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(path.resolve()))
            doc.SaveAs(str(converted.resolve()), FileFormat=16)
            doc.Close()
        finally:
            word.Quit()
        work_path = converted

    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"DOCX parser import failed: {exc}") from exc

    doc = Document(str(work_path))
    blocks: list[dict] = []
    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue
        style_name = para.style.name if para.style else "Normal"
        block_type = "heading" if "Heading" in style_name else "paragraph"
        blocks.append({"type": block_type, "style": style_name, "text": text})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append({"type": "table", "style": "Table", "text": "\n".join(rows)})

    return blocks


__all__ = ["extract_docx_like_blocks"]
