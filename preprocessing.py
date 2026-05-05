from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import os
import re
import site
import shutil
import subprocess
import sys
import sysconfig
import tempfile
import uuid
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path

log = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parent
LOG_DIR = PROJECT_ROOT / "logs"
LOG_FILE = LOG_DIR / "preprocessing.log"


def configure_logging() -> None:
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.StreamHandler(sys.stdout),
            logging.FileHandler(LOG_FILE, encoding="utf-8"),
        ],
    )

SUPPORTED_EXTS = {
    ".pdf",
    ".docx",
    ".doc",
    ".hwp",
    ".hwpx",
    ".csv",
    ".txt",
}
DEFAULT_CHUNK_SIZE = 900
DEFAULT_CHUNK_OVERLAP = 120
DEFAULT_PDF_OCR_MODE = "auto"
DEFAULT_OCR_LANGUAGE = "kor+eng"
DEFAULT_OCR_DPI = 200
MIN_USEFUL_PDF_TEXT_CHARS = 80


def normalize_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\u00a0", " ")
    text = " ".join(text.split())
    return text.strip()


def is_likely_broken_korean_text(text: str) -> bool:
    normalized = normalize_text(text)
    if len(normalized) < MIN_USEFUL_PDF_TEXT_CHARS:
        return False
    letters = re.findall(r"[A-Za-z\uac00-\ud7a3\u4e00-\u9fff]", normalized)
    if not letters:
        return False
    hangul_count = len(re.findall(r"[\uac00-\ud7a3]", normalized))
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", normalized))
    replacement_count = normalized.count("\ufffd") + normalized.count("?")
    return (
        (cjk_count / len(letters) >= 0.15 and hangul_count / len(letters) <= 0.10)
        or replacement_count / max(1, len(normalized)) >= 0.08
        or normalized.count("??") >= 8
    )


def rel_project_path(path: Path, root: Path) -> str:
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError:
        return path.resolve().as_posix()


def make_slug(rel_path: str) -> str:
    return hashlib.md5(rel_path.encode("utf-8")).hexdigest()[:12]


def ensure_output_path(input_file: Path, input_root: Path, output_root: Path) -> Path:
    rel = input_file.resolve().relative_to(input_root.resolve())
    ext_dir = input_file.suffix.lower().lstrip(".") or "unknown"
    return (output_root / ext_dir / rel).with_suffix(".json")


def load_attachment_index(output_json_root: Path, project_root: Path) -> dict[str, dict]:
    index: dict[str, dict] = {}
    if not output_json_root.exists():
        return index

    for jf in output_json_root.rglob("*.json"):
        try:
            doc = json.loads(jf.read_text(encoding="utf-8"))
        except Exception:
            continue

        base_doc_info = {
            "doc_title": doc.get("title", ""),
            "doc_url": doc.get("url", ""),
            "category": doc.get("category", ""),
            "subcategory": doc.get("subcategory", ""),
        }
        for a in doc.get("attachments", []) or []:
            saved_path = a.get("saved_path", "")
            if not saved_path:
                continue
            saved_rel = rel_project_path(project_root / saved_path, project_root)
            index[saved_rel] = {
                **base_doc_info,
                "attachment_name": a.get("name", ""),
                "attachment_url": a.get("url", ""),
                "source_page_url": a.get("source_page_url", ""),
                "source_site": a.get("source_site", ""),
                "downloaded_from_url": a.get("downloaded_from_url", ""),
                "content_type": a.get("content_type", ""),
            }
    return index


def extract_pdf_text_blocks(path: Path) -> list[dict]:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"PDF parser import failed (PyMuPDF/fitz required): {exc}") from exc

    def is_likely_page_number(line: str) -> bool:
        line = line.strip()
        return bool(re.fullmatch(r"-?\s*\d+\s*-?", line))

    def detect_repeated_headers_footers(
        page_lines_list: list[list[str]], min_repeat: int = 3
    ) -> set[str]:
        candidates: list[str] = []
        for lines in page_lines_list:
            if not lines:
                continue
            top = lines[:2]
            bottom = lines[-2:] if len(lines) >= 2 else lines
            for line in top + bottom:
                line = normalize_text(line)
                if line:
                    candidates.append(line)
        counter = Counter(candidates)
        return {line for line, count in counter.items() if count >= min_repeat}

    def should_merge(prev_line: str, curr_line: str) -> bool:
        prev_line = prev_line.strip()
        curr_line = curr_line.strip()
        if not prev_line or not curr_line:
            return False
        if prev_line.endswith((".", "!", "?", ":", ";")):
            return False
        if len(prev_line) < 20:
            return False
        if re.match(r"^[a-z0-9(\[\-]", curr_line):
            return True
        return True

    def merge_lines_into_paragraphs(lines: list[str]) -> list[str]:
        paragraphs: list[str] = []
        buffer: list[str] = []
        for line in lines:
            line = normalize_text(line)
            if not line:
                if buffer:
                    paragraphs.append(" ".join(buffer).strip())
                    buffer = []
                continue
            if is_likely_page_number(line):
                continue
            if not buffer:
                buffer.append(line)
                continue
            prev_line = buffer[-1]
            if should_merge(prev_line, line):
                buffer.append(line)
            else:
                paragraphs.append(" ".join(buffer).strip())
                buffer = [line]
        if buffer:
            paragraphs.append(" ".join(buffer).strip())
        return paragraphs

    def remove_consecutive_duplicate_paragraphs(paragraphs: list[str]) -> list[str]:
        cleaned: list[str] = []
        prev = None
        for para in paragraphs:
            para = normalize_text(para)
            if not para:
                continue
            if para != prev:
                cleaned.append(para)
            prev = para
        return cleaned

    doc = fitz.open(str(path))
    all_page_lines: list[list[str]] = []
    for page in doc:
        text = page.get_text("text")
        lines = [normalize_text(line) for line in text.splitlines()]
        lines = [line for line in lines if line]
        all_page_lines.append(lines)

    repeated_headers_footers = detect_repeated_headers_footers(all_page_lines)
    blocks: list[dict] = []
    for page_num, lines in enumerate(all_page_lines, start=1):
        cleaned_lines: list[str] = []
        for line in lines:
            line = normalize_text(line)
            if not line:
                continue
            if line in repeated_headers_footers:
                continue
            if is_likely_page_number(line):
                continue
            cleaned_lines.append(line)

        paragraphs = merge_lines_into_paragraphs(cleaned_lines)
        paragraphs = remove_consecutive_duplicate_paragraphs(paragraphs)
        for para in paragraphs:
            blocks.append(
                {
                    "type": "paragraph",
                    "style": "Normal",
                    "page": page_num,
                    "text": para,
                }
            )

    doc.close()
    return blocks


def extract_pdf_ocr_blocks(
    path: Path,
    ocr_language: str = DEFAULT_OCR_LANGUAGE,
    ocr_dpi: int = DEFAULT_OCR_DPI,
) -> list[dict]:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"PDF parser import failed (PyMuPDF/fitz required): {exc}") from exc

    tesseract = shutil.which("tesseract")
    if not tesseract:
        log.warning("[OCR-SKIP] tesseract command not found: %s", path)
        return []

    blocks: list[dict] = []
    doc = fitz.open(str(path))
    temp_dir = Path(tempfile.mkdtemp(prefix="campus_rag_ocr_"))
    try:
        zoom = max(72, ocr_dpi) / 72
        matrix = fitz.Matrix(zoom, zoom)
        for page_num, page in enumerate(doc, start=1):
            image_path = temp_dir / f"page-{page_num:04d}.png"
            page.get_pixmap(matrix=matrix, alpha=False).save(str(image_path))
            proc = subprocess.run(
                [tesseract, str(image_path), "stdout", "-l", ocr_language, "--psm", "6"],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
                check=False,
            )
            if proc.returncode != 0:
                err = normalize_text(proc.stderr) or "unknown error"
                log.warning("[OCR-FAIL] %s page %d (%s)", path, page_num, err)
                continue
            text = normalize_text(proc.stdout)
            if text:
                blocks.append(
                    {
                        "type": "ocr_paragraph",
                        "style": "Tesseract",
                        "page": page_num,
                        "text": text,
                    }
                )
    finally:
        doc.close()
        shutil.rmtree(temp_dir, ignore_errors=True)
    return blocks


def extract_pdf_blocks(
    path: Path,
    ocr_mode: str = DEFAULT_PDF_OCR_MODE,
    ocr_language: str = DEFAULT_OCR_LANGUAGE,
    ocr_dpi: int = DEFAULT_OCR_DPI,
) -> list[dict]:
    blocks = extract_pdf_text_blocks(path)
    text = "\n".join(normalize_text(b.get("text", "")) for b in blocks)
    needs_ocr = (
        ocr_mode == "always"
        or (
            ocr_mode == "auto"
            and (
                len(normalize_text(text)) < MIN_USEFUL_PDF_TEXT_CHARS
                or is_likely_broken_korean_text(text)
            )
        )
    )
    if not needs_ocr:
        return blocks
    reason = "empty/short text" if len(normalize_text(text)) < MIN_USEFUL_PDF_TEXT_CHARS else "broken text"
    log.info("[OCR] %s (%s)", path, reason)
    ocr_blocks = extract_pdf_ocr_blocks(path, ocr_language=ocr_language, ocr_dpi=ocr_dpi)
    return ocr_blocks or blocks


def extract_docx_like_blocks(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    work_path = path
    if suffix == ".doc":
        try:
            import win32com.client  # type: ignore
        except Exception as exc:
            raise RuntimeError(f".doc conversion requires pywin32: {exc}") from exc

        converted = path.with_suffix(".docx")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(path.resolve()))
            doc.SaveAs(str(converted.resolve()), FileFormat=16)
            doc.Close()
        finally:
            word.Quit()
        work_path = converted

    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"DOCX parser import failed: {exc}") from exc

    doc = Document(str(work_path))
    blocks: list[dict] = []
    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue
        style_name = para.style.name if para.style else "Normal"
        block_type = "heading" if "Heading" in style_name else "paragraph"
        blocks.append({"type": block_type, "style": style_name, "text": text})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append({"type": "table", "style": "Table", "text": "\n".join(rows)})

    return blocks


def extract_csv_blocks(path: Path) -> list[dict]:
    blocks: list[dict] = []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(f, dialect)
        for row_idx, row in enumerate(reader, start=1):
            cells = [normalize_text(c) for c in row if normalize_text(c)]
            if not cells:
                continue
            blocks.append(
                {
                    "type": "table_row",
                    "style": "CSV",
                    "row": row_idx,
                    "text": " | ".join(cells),
                }
            )
    return blocks


def extract_txt_blocks(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = [normalize_text(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    blocks = []
    for idx, line in enumerate(lines, start=1):
        blocks.append(
            {
                "type": "paragraph",
                "style": "Text",
                "line": idx,
                "text": line,
            }
        )
    return blocks


def extract_html_blocks(path: Path) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
        from bs4.element import Tag
    except Exception as exc:
        raise RuntimeError(f"HTML parser import failed (beautifulsoup4 required): {exc}") from exc

    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe", "form"]):
        tag.decompose()

    def attr_text(tag: Tag) -> str:
        attrs: list[str] = []
        tag_id = tag.get("id")
        if tag_id:
            attrs.append(str(tag_id))
        classes = tag.get("class") or []
        attrs.extend(str(cls) for cls in classes)
        return " ".join(attrs).lower()

    def is_noise(tag: Tag) -> bool:
        if tag.name in {"nav", "header", "footer", "aside"}:
            return True
        attrs = attr_text(tag)
        noise_keywords = (
            "nav",
            "gnb",
            "lnb",
            "snb",
            "menu",
            "header",
            "footer",
            "breadcrumb",
            "quick",
            "search",
            "banner",
            "popup",
            "comment",
            "reply",
            "pagination",
            "print",
            "share",
            "zoom",
            "btn",
            "button",
        )
        return any(keyword in attrs for keyword in noise_keywords)

    def text_without_tables(tag: Tag) -> str:
        parts: list[str] = []
        for child in tag.children:
            if isinstance(child, str):
                parts.append(child)
                continue
            if not isinstance(child, Tag):
                continue
            if child.name == "table":
                continue
            parts.append(text_without_tables(child))
        return normalize_text(" ".join(parts))

    def element_score(tag: Tag) -> int:
        text = normalize_text(tag.get_text(" ", strip=True))
        if len(text) < 80:
            return -1
        attrs = attr_text(tag)
        score = len(text)
        if any(keyword in attrs for keyword in ("content", "cont", "board", "view", "article", "body", "main", "read", "detail")):
            score += 400
        if any(keyword in attrs for keyword in ("nav", "menu", "header", "footer", "banner", "search", "quick")):
            score -= 300
        return score

    body = soup.body or soup
    candidates: list[Tag] = [body]
    for tag in body.find_all(["main", "article", "section", "div", "td"]):
        if is_noise(tag):
            continue
        candidates.append(tag)

    root = max(candidates, key=element_score)

    blocks: list[dict] = []
    seen_texts: set[str] = set()
    ignored_texts = {
        "확대",
        "축소",
        "프린트",
        "인쇄",
        "공유",
        "목록",
        "다음글",
        "이전글",
    }

    def append_block(block_type: str, style: str, text: str) -> None:
        text = normalize_text(text)
        if not text or text in seen_texts or text in ignored_texts:
            return
        seen_texts.add(text)
        blocks.append({"type": block_type, "style": style, "text": text})

    title = normalize_text((soup.title.string if soup.title and soup.title.string else ""))
    if title:
        append_block("heading", "HTMLTitle", title)

    def table_to_text(table_tag: Tag) -> str:
        rows: list[str] = []
        for tr in table_tag.find_all("tr"):
            cells: list[str] = []
            for cell in tr.find_all(["th", "td"], recursive=False):
                cell_text = text_without_tables(cell)
                if cell_text:
                    cells.append(cell_text)
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def walk(tag: Tag) -> None:
        if is_noise(tag):
            return

        if tag.name == "table":
            table_text = table_to_text(tag)
            if table_text:
                append_block("table", "HTMLTable", table_text)
            return

        if tag.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            append_block("heading", tag.name.upper(), text_without_tables(tag))
        elif tag.name == "p":
            append_block("paragraph", "HTMLParagraph", text_without_tables(tag))
        elif tag.name == "li":
            append_block("list_item", "HTMLListItem", text_without_tables(tag))
        elif tag.name in {"blockquote", "pre"}:
            append_block("paragraph", tag.name.upper(), text_without_tables(tag))

        for child in tag.children:
            if isinstance(child, Tag):
                walk(child)

    walk(root)

    if blocks:
        return blocks

    fallback_lines = [
        normalize_text(line)
        for line in root.get_text("\n", strip=True).splitlines()
        if normalize_text(line)
    ]
    for line in fallback_lines:
        append_block("paragraph", "HTMLText", line)
    return blocks


def extract_hwpx_blocks(path: Path) -> list[dict]:
    blocks: list[dict] = []
    with zipfile.ZipFile(path) as zf:
        xml_names = [n for n in zf.namelist() if n.lower().endswith(".xml")]
        xml_names.sort()
        for name in xml_names:
            data = zf.read(name)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            texts = []
            for elem in root.iter():
                if elem.text and normalize_text(elem.text):
                    texts.append(normalize_text(elem.text))
            if not texts:
                continue
            blocks.append(
                {
                    "type": "xml_text",
                    "style": "HWPX",
                    "section": name,
                    "text": " ".join(texts),
                }
            )
    return blocks


def extract_hwp_blocks(path: Path) -> list[dict]:
    def format_hwp_runtime_error(stderr: str) -> str:
        err = stderr.strip() or "unknown error"
        if "No module named 'six'" in err or 'No module named "six"' in err:
            return (
                "pyhwp runtime dependency 'six' is missing in the interpreter used by "
                "hwp5txt/hwp5html. Install it in the same environment, for example: "
                f'"{sys.executable}" -m pip install six'
            )
        return err

    def resolve_hwp_command(command_name: str) -> str | None:
        resolved = shutil.which(command_name)
        if resolved:
            return resolved

        if os.name != "nt":
            return None

        exe_name = f"{command_name}.exe"
        candidate_paths = [
            Path(sys.executable).resolve().parent / "Scripts" / exe_name,
            Path(sys.executable).resolve().parent / exe_name,
            Path(sysconfig.get_path("scripts")) / exe_name,
            Path(site.getuserbase()) / "Python313" / "Scripts" / exe_name,
        ]
        for candidate in candidate_paths:
            if candidate.exists():
                return str(candidate)
        return None

    def extract_hwp_blocks_from_html(path: Path) -> list[dict]:
        hwp5html = resolve_hwp_command("hwp5html")
        if not hwp5html:
            return []

        tmp_base = Path(tempfile.gettempdir()) / "campus_rag_hwp5html"
        tmp_base.mkdir(parents=True, exist_ok=True)
        token = uuid.uuid4().hex
        copied_hwp = tmp_base / f"{token}.hwp"
        html_file = tmp_base / f"{token}.xhtml"
        try:
            shutil.copy2(path, copied_hwp)
            proc = subprocess.run(
                [hwp5html, "--html", "--output", str(html_file), str(copied_hwp)],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
                check=False,
            )
            if proc.returncode != 0:
                log.debug("hwp5html failed for %s: %s", path, format_hwp_runtime_error(proc.stderr))
                return []

            if not html_file.exists():
                return []

            try:
                root = ET.fromstring(html_file.read_text(encoding="utf-8", errors="ignore"))
            except ET.ParseError:
                return []

            def tag_name(elem: ET.Element) -> str:
                return elem.tag.rsplit("}", 1)[-1] if "}" in elem.tag else elem.tag

            def collect_text(elem: ET.Element) -> str:
                return normalize_text("".join(elem.itertext()))

            def table_to_text(table_elem: ET.Element) -> str:
                rows: list[str] = []
                for tr in table_elem.iter():
                    if tag_name(tr) != "tr":
                        continue
                    cells: list[str] = []
                    for child in list(tr):
                        if tag_name(child) not in {"td", "th"}:
                            continue
                        cell_text = collect_text(child)
                        if cell_text:
                            cells.append(cell_text)
                    if cells:
                        rows.append(" | ".join(cells))
                return "\n".join(rows)

            body = None
            for elem in root.iter():
                if tag_name(elem) == "body":
                    body = elem
                    break
            if body is None:
                return []

            blocks: list[dict] = []

            def collect_text_without_tables(elem: ET.Element) -> str:
                parts: list[str] = []
                if elem.text:
                    parts.append(elem.text)
                for child in list(elem):
                    if tag_name(child) != "table":
                        parts.append(collect_text_without_tables(child))
                    if child.tail:
                        parts.append(child.tail)
                return "".join(parts)

            def walk(elem: ET.Element) -> None:
                name = tag_name(elem)
                if name == "table":
                    table_text = table_to_text(elem)
                    if table_text:
                        blocks.append(
                            {
                                "type": "table",
                                "style": "HWP",
                                "text": table_text,
                            }
                        )
                    return
                if name == "p":
                    text = normalize_text(collect_text_without_tables(elem))
                    if text:
                        blocks.append(
                            {
                                "type": "paragraph",
                                "style": "HWP",
                                "text": text,
                            }
                        )
                for child in list(elem):
                    walk(child)

            walk(body)
            return blocks
        finally:
            try:
                copied_hwp.unlink(missing_ok=True)
            except Exception:
                pass
            try:
                html_file.unlink(missing_ok=True)
            except Exception:
                pass

    hwp5txt = resolve_hwp_command("hwp5txt")

    if not hwp5txt:
        raise RuntimeError(
            "hwp5txt command not found in PATH/current Python environment. "
            "Install pyhwp in the same interpreter and add its Scripts directory to PATH."
        )

    html_blocks = extract_hwp_blocks_from_html(path)
    if html_blocks:
        return html_blocks

    proc = subprocess.run(
        [hwp5txt, str(path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="ignore",
        check=False,
    )
    if proc.returncode != 0:
        raise RuntimeError(f"hwp5txt failed: {format_hwp_runtime_error(proc.stderr)}")

    lines = [normalize_text(line) for line in proc.stdout.splitlines()]
    lines = [line for line in lines if line]

    blocks = []
    for idx, line in enumerate(lines, start=1):
        blocks.append(
            {
                "type": "paragraph",
                "style": "HWP",
                "line": idx,
                "text": line,
            }
        )
    return blocks


def extract_blocks(
    path: Path,
    pdf_ocr_mode: str = DEFAULT_PDF_OCR_MODE,
    ocr_language: str = DEFAULT_OCR_LANGUAGE,
    ocr_dpi: int = DEFAULT_OCR_DPI,
) -> list[dict]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_blocks(
            path,
            ocr_mode=pdf_ocr_mode,
            ocr_language=ocr_language,
            ocr_dpi=ocr_dpi,
        )
    if ext in {".docx", ".doc"}:
        return extract_docx_like_blocks(path)
    if ext in {".html", ".htm"}:
        return extract_html_blocks(path)
    if ext == ".csv":
        return extract_csv_blocks(path)
    if ext == ".txt":
        return extract_txt_blocks(path)
    if ext == ".hwpx":
        return extract_hwpx_blocks(path)    
    if ext == ".hwp":
        return extract_hwp_blocks(path)
    raise ValueError(f"unsupported extension: {ext}")


def build_metadata_fallback_blocks(
    input_file: Path,
    rel: str,
    rel_in_input: str,
    provenance: dict,
) -> list[dict]:
    fields = [
        ("문서 파일명", input_file.name),
        ("문서 경로", rel_in_input),
        ("문서 제목", provenance.get("doc_title", "")),
        ("게시글 URL", provenance.get("doc_url", "")),
        ("카테고리", provenance.get("category", "")),
        ("하위 카테고리", provenance.get("subcategory", "")),
        ("첨부파일명", provenance.get("attachment_name", "")),
        ("첨부파일 URL", provenance.get("attachment_url", "")),
        ("출처 페이지", provenance.get("source_page_url", "")),
        ("출처 사이트", provenance.get("source_site", "")),
        ("저장 경로", rel),
    ]
    lines = [f"{label}: {value}" for label, value in fields if normalize_text(str(value))]
    if not lines:
        return []
    return [{"type": "metadata_fallback", "style": "Metadata", "text": "\n".join(lines)}]


def chunk_blocks(
    blocks: list[dict],
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[dict]:
    texts = [normalize_text(b.get("text", "")) for b in blocks]
    texts = [t for t in texts if t]
    if not texts:
        return []

    chunks: list[dict] = []
    current: list[str] = []
    current_len = 0
    idx = 1

    for text in texts:
        add_len = len(text) + (1 if current else 0)
        if current and (current_len + add_len > chunk_size):
            chunk_text = "\n".join(current)
            chunks.append(
                {
                    "chunk_id": idx,
                    "text": chunk_text,
                    "num_lines": len(current),
                    "num_chars": len(chunk_text),
                }
            )
            idx += 1

            if overlap > 0:
                tail = chunk_text[-overlap:]
                current = [tail] if tail.strip() else []
                current_len = len(tail) if current else 0
            else:
                current = []
                current_len = 0

        current.append(text)
        current_len += len(text) + (1 if len(current) > 1 else 0)

    if current:
        chunk_text = "\n".join(current)
        chunks.append(
            {
                "chunk_id": idx,
                "text": chunk_text,
                "num_lines": len(current),
                "num_chars": len(chunk_text),
            }
        )
    return chunks


def save_preprocessed_file(
    input_file: Path,
    input_root: Path,
    output_root: Path,
    project_root: Path,
    attachment_index: dict[str, dict],
    chunk_size: int,
    chunk_overlap: int,
    pdf_ocr_mode: str,
    ocr_language: str,
    ocr_dpi: int,
) -> tuple[bool, str]:
    rel = rel_project_path(input_file, project_root)
    ext = input_file.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        return False, "unsupported"

    rel_in_input = rel_project_path(input_file, input_root)
    provenance = attachment_index.get(rel, {})
    blocks = extract_blocks(
        input_file,
        pdf_ocr_mode=pdf_ocr_mode,
        ocr_language=ocr_language,
        ocr_dpi=ocr_dpi,
    )
    blocks = [b for b in blocks if normalize_text(b.get("text", ""))]
    chunks = chunk_blocks(blocks, chunk_size=chunk_size, overlap=chunk_overlap)
    used_metadata_fallback = False
    if not chunks:
        fallback_blocks = build_metadata_fallback_blocks(input_file, rel, rel_in_input, provenance)
        fallback_chunks = chunk_blocks(fallback_blocks, chunk_size=chunk_size, overlap=0)
        if fallback_chunks:
            blocks = fallback_blocks
            chunks = fallback_chunks
            used_metadata_fallback = True

    out_path = ensure_output_path(input_file, input_root, output_root)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    result = {
        "slug": make_slug(rel),
        "source_file": input_file.name,
        "source_path": rel,
        "source_relative_to_input": rel_in_input,
        "source_ext": ext,
        "processed_at": datetime.now().isoformat(),
        "num_blocks": len(blocks),
        "total_chars": sum(len(b["text"]) for b in blocks),
        "num_chunks": len(chunks),
        "used_metadata_fallback": used_metadata_fallback,
        "provenance": provenance,
        "blocks": blocks,
        "chunks": chunks,
    }

    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return True, out_path.as_posix()


def iter_target_files(input_root: Path) -> list[Path]:
    if not input_root.exists():
        return []
    files = []
    for p in input_root.rglob("*"):
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
            files.append(p)
    return sorted(files)


def run_batch(
    input_root: Path,
    output_root: Path,
    output_json_root: Path,
    project_root: Path,
    dry_run: bool,
    chunk_size: int,
    chunk_overlap: int,
    pdf_ocr_mode: str,
    ocr_language: str,
    ocr_dpi: int,
) -> None:
    attachment_index = load_attachment_index(output_json_root, project_root)
    files = iter_target_files(input_root)
    if not files:
        log.info("처리 대상 파일이 없습니다: %s", input_root)
        return

    log.info("처리 대상: %d개", len(files))
    ok = 0
    skipped = 0
    failed = 0

    for file_path in files:
        rel = rel_project_path(file_path, project_root)
        if dry_run:
            log.info("[DRY-RUN] %s", rel)
            continue
        try:
            saved, info = save_preprocessed_file(
                input_file=file_path,
                input_root=input_root,
                output_root=output_root,
                project_root=project_root,
                attachment_index=attachment_index,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                pdf_ocr_mode=pdf_ocr_mode,
                ocr_language=ocr_language,
                ocr_dpi=ocr_dpi,
            )
            if saved:
                ok += 1
                log.info("[OK] %s -> %s", rel, info)
            else:
                skipped += 1
                log.info("[SKIP] %s (%s)", rel, info)
        except Exception as exc:
            failed += 1
            log.warning("[FAIL] %s (%s)", rel, exc)

    if dry_run:
        log.info("DRY-RUN 완료")
    else:
        log.info("완료: 성공=%d, 건너뜀=%d, 실패=%d", ok, skipped, failed)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="FILES/output/files 첨부파일을 RAG용 전처리 JSON으로 변환"
    )
    parser.add_argument(
        "--input-root",
        type=Path,
        default=Path("FILES/output/files"),
        help="원본 첨부파일 루트 경로",
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path("FILES/preprocessed"),
        help="전처리 JSON 출력 루트 경로",
    )
    parser.add_argument(
        "--output-json-root",
        type=Path,
        default=Path("FILES/output/json"),
        help="크롤링 본문 JSON 루트 (첨부 provenance 조인용)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="실제 저장 없이 대상 파일만 확인",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=DEFAULT_CHUNK_SIZE,
        help="RAG 임베딩용 청크 최대 문자 수",
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=DEFAULT_CHUNK_OVERLAP,
        help="연속 청크 간 문자 오버랩 길이",
    )
    parser.add_argument(
        "--pdf-ocr",
        choices=["auto", "never", "always"],
        default=DEFAULT_PDF_OCR_MODE,
        help="PDF OCR fallback 사용 방식",
    )
    parser.add_argument(
        "--ocr-language",
        default=DEFAULT_OCR_LANGUAGE,
        help="Tesseract OCR 언어 설정",
    )
    parser.add_argument(
        "--ocr-dpi",
        type=int,
        default=DEFAULT_OCR_DPI,
        help="PDF 페이지를 OCR 이미지로 렌더링할 DPI",
    )
    args = parser.parse_args()

    configure_logging()
    project_root = Path.cwd()
    run_batch(
        input_root=args.input_root,
        output_root=args.output_root,
        output_json_root=args.output_json_root,
        project_root=project_root,
        dry_run=args.dry_run,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        pdf_ocr_mode=args.pdf_ocr,
        ocr_language=args.ocr_language,
        ocr_dpi=args.ocr_dpi,
    )


if __name__ == "__main__":
    main()
